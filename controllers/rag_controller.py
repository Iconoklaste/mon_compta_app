# c:\wamp\www\mon_compta_app\controllers\rag_controller.py
import os
import logging
import numpy as np
from flask import Blueprint, request, jsonify, current_app, session, flash, send_from_directory, url_for, redirect
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from controllers.db_manager import db
from models import Projet, User, RagDocument # Assurez-vous que RagDocument est bien importé

# --- Bibliothèques pour l'extraction de texte ---
import PyPDF2
import docx # python-docx
import re # Ajouté pour le nettoyage de texte

# --- Client Mistral ---
from mistralai import Mistral

# Initialiser le logger pour ce contrôleur
logger = logging.getLogger(__name__)

# Création du Blueprint 'rag_bp'
rag_bp = Blueprint('rag', __name__, url_prefix='/rag')

# --- Constantes ---
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md'} # Ajout de .md par exemple
MISTRAL_EMBEDDING_MODEL = "mistral-embed"

# --- Fonctions Utilitaires d'extraction (inchangées par rapport à votre version) ---
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_stream):
    text = ""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text() or ""
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte PDF: {e}", exc_info=True)
        return ""
    return text

def extract_text_from_docx(file_stream):
    text = ""
    try:
        document = docx.Document(file_stream)
        for para in document.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOCX: {e}", exc_info=True)
        return ""
    return text

def extract_text_from_txt(file_stream):
    try:
        return file_stream.read().decode('utf-8', errors='replace')
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte TXT: {e}", exc_info=True)
        return ""

def chunk_text(text, chunk_size=500, chunk_overlap=50, separators=None):
    if not text or not isinstance(text, str):
        logger.warning("chunk_text a reçu un texte vide ou invalide.")
        return []
    if separators is None:
        separators = ["\n\n", "\n", ". ", "! ", "? ", " ", ""]
    final_chunks = []
    if len(text) <= chunk_size:
        stripped_text = text.strip()
        if stripped_text: return [stripped_text]
        return []
    current_separator = ""
    for sep in separators:
        if sep in text:
            current_separator = sep
            break
    splits = text.split(current_separator) if current_separator else [text]
    intermediate_documents = []
    current_doc_content = ""
    for i, split_part in enumerate(splits):
        part_to_add = split_part
        if current_separator and i < len(splits) - 1 :
            part_to_add += current_separator
        if len(current_doc_content) + len(part_to_add) > chunk_size and current_doc_content:
            if current_doc_content.strip():
                intermediate_documents.append(current_doc_content.strip())
            current_doc_content = part_to_add
        else:
            current_doc_content += part_to_add
    if current_doc_content.strip():
        intermediate_documents.append(current_doc_content.strip())
    if not intermediate_documents and text.strip():
        intermediate_documents = [text.strip()]
    for doc_item in intermediate_documents:
        if not doc_item.strip(): continue
        if len(doc_item) <= chunk_size:
            final_chunks.append(doc_item)
            continue
        start = 0
        doc_len = len(doc_item)
        while start < doc_len:
            end = start + chunk_size
            chunk = doc_item[start:end].strip()
            if chunk: final_chunks.append(chunk)
            next_start = start + chunk_size - chunk_overlap
            if next_start <= start or end >= doc_len:
                start += chunk_size
            else:
                start = next_start
    seen_chunks = set()
    unique_final_chunks = []
    for c in final_chunks:
        if c and c not in seen_chunks:
            unique_final_chunks.append(c)
            seen_chunks.add(c)
    logger.info(f"Chunking terminé. {len(unique_final_chunks)} chunks finaux générés.")
    return unique_final_chunks

def get_mistral_embeddings(texts, api_key):
    if not texts: return []
    try:
        client = Mistral(api_key=api_key)
        embeddings_batch_response = client.embeddings.create(
            model=MISTRAL_EMBEDDING_MODEL,
            inputs=texts
        )
        return [data.embedding for data in embeddings_batch_response.data]
    except Exception as e:
        logger.error(f"Erreur lors de l'appel à l'API d'embedding Mistral: {e}", exc_info=True)
        return []

def get_rag_filepath(projet_id):
    rag_storage_dir = current_app.config.get('RAG_STORAGE_PATH')
    if not rag_storage_dir:
        logger.error("RAG_STORAGE_PATH n'est pas configuré.")
        raise ValueError("Configuration RAG_STORAGE_PATH manquante.")
    return os.path.join(rag_storage_dir, f"rag_projet_{projet_id}.npz")

def get_original_document_save_path(projet_id, filename):
    upload_folder = current_app.config.get('RAG_UPLOAD_FOLDER')
    if not upload_folder:
        logger.error("RAG_UPLOAD_FOLDER n'est pas configuré.")
        raise ValueError("Configuration RAG_UPLOAD_FOLDER manquante.")
    project_specific_upload_folder = os.path.join(upload_folder, str(projet_id))
    os.makedirs(project_specific_upload_folder, exist_ok=True)
    return os.path.join(project_specific_upload_folder, filename)

# --- Route d'Indexation (Stratégie Cumulative) ---
@rag_bp.route('/projets/<int:projet_id>/indexer_document', methods=['POST'])
@login_required
def indexer_document_route(projet_id):
    projet = Projet.query.filter_by(id=projet_id, organisation_id=current_user.organisation_id).first_or_404()
    if not projet: # Should be handled by first_or_404
        return jsonify({"success": False, "error": "Projet non trouvé ou accès non autorisé."}), 403

    if 'document' not in request.files:
        return jsonify({"success": False, "error": "Aucun fichier fourni."}), 400
    file = request.files['document']
    if file.filename == '':
        return jsonify({"success": False, "error": "Aucun fichier sélectionné."}), 400
    if not allowed_file(file.filename):
        return jsonify({"success": False, "error": f"Type de fichier non autorisé. Acceptés: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

    original_filename = secure_filename(file.filename)
    logger.info(f"Début de l'indexation cumulative du document '{original_filename}' pour projet ID {projet_id}.")
    original_file_save_path = None
    new_rag_doc_entry_id = None # Pour le cleanup en cas d'erreur après création DB

    try:
        # 0. Vérifier si un document avec le même nom existe déjà pour ce projet.
        #    Si oui, on le supprime d'abord (logique de mise à jour).
        existing_doc_same_name = RagDocument.query.filter_by(projet_id=projet_id, original_filename=original_filename).first()
        if existing_doc_same_name:
            logger.info(f"Un document nommé '{original_filename}' existe déjà (ID: {existing_doc_same_name.id}). Il sera remplacé.")
            # Appeler la logique de suppression pour cet ancien document
            # Note: cette fonction interne ne retourne pas de réponse HTTP mais lève des exceptions
            _internal_delete_rag_document_logic(projet_id, existing_doc_same_name.id, existing_doc_same_name.original_filename)
            logger.info(f"Ancien document '{original_filename}' (ID: {existing_doc_same_name.id}) supprimé avant ré-indexation.")


        # 1. Sauvegarder le nouveau fichier original
        original_file_save_path = get_original_document_save_path(projet_id, original_filename)
        file.save(original_file_save_path)
        logger.info(f"Fichier original '{original_filename}' sauvegardé dans '{original_file_save_path}'.")
        file.seek(0) # Rembobiner pour l'extraction

        # 2. Enregistrer l'entrée RagDocument en BDD pour obtenir un ID
        new_rag_doc_entry = RagDocument(
            projet_id=projet_id,
            original_filename=original_filename,
            status='En cours d\'indexation'
        )
        db.session.add(new_rag_doc_entry)
        db.session.commit() # Commit pour obtenir l'ID
        new_rag_doc_entry_id = new_rag_doc_entry.id
        logger.info(f"Entrée RagDocument créée (ID: {new_rag_doc_entry_id}) pour '{original_filename}'.")

        # 3. Extraction de texte
        text_content = ""
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        if file_extension == 'pdf': text_content = extract_text_from_pdf(file.stream)
        elif file_extension == 'docx': text_content = extract_text_from_docx(file.stream)
        elif file_extension == 'txt': text_content = extract_text_from_txt(file.stream)
        elif file_extension == 'md': text_content = extract_text_from_txt(file.stream) # Traiter .md comme .txt

        if not text_content or not text_content.strip():
            new_rag_doc_entry.status = 'Erreur (contenu vide)'
            db.session.commit()
            return jsonify({"success": False, "error": "Impossible d'extraire le contenu ou document vide."}), 400
        
        text_content = re.sub(r'\s+', ' ', text_content).strip()
        logger.info(f"Texte extrait de '{original_filename}' (longueur: {len(text_content)}).")

        # 4. Chunking
        text_chunks = chunk_text(text_content)
        if not text_chunks:
            new_rag_doc_entry.status = 'Erreur (chunking impossible)'
            db.session.commit()
            return jsonify({"success": False, "error": "Document non découpable en morceaux."}), 400
        logger.info(f"{len(text_chunks)} chunks créés pour '{original_filename}'.")

        # 5. Génération des Embeddings
        api_key = os.environ.get("MISTRAL_API_KEY")
        if not api_key:
            new_rag_doc_entry.status = 'Erreur (API Key)'
            db.session.commit()
            return jsonify({"success": False, "error": "Erreur config serveur (API Key manquante)."}), 500

        new_doc_embeddings_list = get_mistral_embeddings(text_chunks, api_key)
        if not new_doc_embeddings_list or len(new_doc_embeddings_list) != len(text_chunks):
            new_rag_doc_entry.status = 'Erreur (embeddings)'
            db.session.commit()
            return jsonify({"success": False, "error": "Erreur génération embeddings."}), 500
        new_doc_embeddings_np = np.array(new_doc_embeddings_list, dtype=np.float32)
        logger.info(f"{len(new_doc_embeddings_list)} embeddings générés pour '{original_filename}'.")

        # 6. Chargement de l'index RAG existant et fusion
        rag_file_path = get_rag_filepath(projet_id)
        
        all_texts_list = []
        all_metadata_list = []
        # S'assurer que existing_embeddings_np a la bonne shape même si vide
        # La dimension (1024) est spécifique à mistral-embed, à rendre configurable si besoin
        embedding_dim = new_doc_embeddings_np.shape[1] if new_doc_embeddings_np.ndim > 1 and new_doc_embeddings_np.size > 0 else 1024
        existing_embeddings_np = np.array([], dtype=np.float32).reshape(0, embedding_dim)

        if os.path.exists(rag_file_path):
            logger.info(f"Chargement du fichier RAG existant : {rag_file_path}")
            try:
                with np.load(rag_file_path, allow_pickle=True) as data:
                    if 'texts' in data and 'embeddings' in data and 'metadata' in data:
                        all_texts_list = list(data['texts'])
                        existing_embeddings_np = data['embeddings']
                        all_metadata_list = list(data['metadata'])
                        if existing_embeddings_np.ndim == 1 and len(all_texts_list) == 1 : # Cas d'un seul embedding mal shapé
                            existing_embeddings_np = existing_embeddings_np.reshape(1, -1)
                        logger.info(f"{len(all_texts_list)} textes, {existing_embeddings_np.shape[0]} embeddings existants chargés.")
                    else:
                        logger.warning(f"Fichier {rag_file_path} mal formaté ou incomplet. Il sera écrasé/réinitialisé.")
            except Exception as e_load:
                logger.error(f"Erreur chargement {rag_file_path}: {e_load}. Fichier sera écrasé/réinitialisé.", exc_info=True)
                # Réinitialiser les listes en cas d'erreur de chargement
                all_texts_list = []
                all_metadata_list = []
                existing_embeddings_np = np.array([], dtype=np.float32).reshape(0, embedding_dim)
        
        # Préparer les métadonnées pour les nouveaux chunks
        new_doc_metadata_list = [
            {"source_filename": original_filename, 
             "original_path": original_file_save_path, # Peut être utile pour débogage
             "rag_document_db_id": new_rag_doc_entry_id # LIEN VERS L'ENTRÉE DB
            } for _ in text_chunks
        ]

        # Ajouter les nouveaux chunks, embeddings et métadonnées
        all_texts_list.extend(text_chunks)
        all_metadata_list.extend(new_doc_metadata_list)

        if existing_embeddings_np.size == 0: # Si l'index était vide ou nouveau
            all_embeddings_np_combined = new_doc_embeddings_np
        elif existing_embeddings_np.shape[1] != new_doc_embeddings_np.shape[1]:
            new_rag_doc_entry.status = 'Erreur (dim embeddings)'
            db.session.commit()
            logger.error(f"Incompatibilité de dimension des embeddings. Existants: {existing_embeddings_np.shape[1]}, Nouveaux: {new_doc_embeddings_np.shape[1]}")
            return jsonify({"success": False, "error": "Incompatibilité de dimension des embeddings."}), 500
        else:
            all_embeddings_np_combined = np.vstack([existing_embeddings_np, new_doc_embeddings_np])
        
        logger.info(f"Index combiné: {len(all_texts_list)} textes, {all_embeddings_np_combined.shape[0]} embeddings.")

        # 7. Sauvegarde de l'index RAG combiné
        try:
            os.makedirs(os.path.dirname(rag_file_path), exist_ok=True)
            np.savez_compressed(rag_file_path,
                                texts=np.array(all_texts_list, dtype=object),
                                embeddings=all_embeddings_np_combined,
                                metadata=np.array(all_metadata_list, dtype=object))
            logger.info(f"Données RAG combinées sauvegardées dans {rag_file_path}.")
        except Exception as e_save_npz:
            new_rag_doc_entry.status = 'Erreur (sauvegarde NPZ)'
            db.session.commit()
            logger.error(f"Erreur sauvegarde RAG {rag_file_path}: {e_save_npz}", exc_info=True)
            return jsonify({"success": False, "error": "Erreur sauvegarde données d'indexation."}), 500

        # Mettre à jour le statut du RagDocument en BDD
        new_rag_doc_entry.status = 'Indexé'
        db.session.commit()

        return jsonify({
            "success": True,
            "message": f"Document '{original_filename}' indexé et ajouté avec succès au projet {projet_id}.",
            "document_id": new_rag_doc_entry_id,
            "chunks_added_from_this_doc": len(text_chunks),
            "total_chunks_in_project_index": len(all_texts_list)
        }), 200

    except ValueError as ve: # Erreurs de configuration
        logger.error(f"Erreur de valeur (config) lors de l'indexation pour projet {projet_id}: {ve}", exc_info=True)
        _cleanup_failed_indexation(original_file_save_path, new_rag_doc_entry_id)
        return jsonify({"success": False, "error": str(ve)}), 400
    except OSError as oe: # Erreurs OS (ex: suppression .npz)
        logger.error(f"Erreur OS lors de l'indexation pour projet {projet_id}: {oe}", exc_info=True)
        _cleanup_failed_indexation(original_file_save_path, new_rag_doc_entry_id)
        return jsonify({"success": False, "error": f"Erreur système: {str(oe)}"}), 500
    except Exception as e:
        logger.error(f"Erreur inattendue lors de l'indexation de '{original_filename}' pour projet {projet_id}: {e}", exc_info=True)
        _cleanup_failed_indexation(original_file_save_path, new_rag_doc_entry_id)
        return jsonify({"success": False, "error": "Erreur serveur lors de l'indexation."}), 500

def _cleanup_failed_indexation(original_file_path, rag_doc_db_id):
    """Nettoie les artefacts si l'indexation échoue après certaines étapes."""
    if original_file_path and os.path.exists(original_file_path):
        try: os.remove(original_file_path); logger.info(f"Nettoyage: Fichier original {original_file_path} supprimé.")
        except: logger.error(f"Nettoyage: Erreur suppression fichier original {original_file_path}.")
    if rag_doc_db_id:
        doc_to_delete = RagDocument.query.get(rag_doc_db_id)
        if doc_to_delete:
            try:
                # Si le statut n'est pas 'Indexé', on peut le supprimer.
                # S'il est 'Indexé' mais que l'erreur est survenue après, la logique de suppression de l'index .npz est plus complexe.
                # Pour l'instant, on se contente de supprimer l'entrée DB si elle n'est pas marquée comme pleinement indexée.
                if doc_to_delete.status != 'Indexé':
                    db.session.delete(doc_to_delete)
                    db.session.commit()
                    logger.info(f"Nettoyage: Entrée RagDocument ID {rag_doc_db_id} supprimée.")
                else: # Si déjà marqué 'Indexé', mais erreur après, c'est un cas limite.
                      # On pourrait changer son statut en 'Erreur (post-indexation)'
                    doc_to_delete.status = 'Erreur (post-indexation)'
                    db.session.commit()
                    logger.warning(f"Nettoyage: Entrée RagDocument ID {rag_doc_db_id} était 'Indexé' mais une erreur est survenue après. Statut mis à 'Erreur (post-indexation)'.")

            except Exception as e_db_clean:
                db.session.rollback()
                logger.error(f"Nettoyage: Erreur suppression entrée RagDocument ID {rag_doc_db_id}: {e_db_clean}")

# --- Route pour Lister les Documents RAG d'un Projet ---
@rag_bp.route('/projets/<int:projet_id>/documents', methods=['GET'])
@login_required
def list_rag_documents_route(projet_id):
    projet = Projet.query.filter_by(id=projet_id, organisation_id=current_user.organisation_id).first_or_404()
    if not projet: # Should be handled by first_or_404
        return jsonify(success=False, error="Projet non trouvé ou accès non autorisé."), 403
    documents = RagDocument.query.filter_by(projet_id=projet_id).order_by(RagDocument.indexed_at.desc()).all()
    return jsonify(success=True, documents=[doc.to_dict() for doc in documents])

# --- Route pour Supprimer un Document RAG (Stratégie Cumulative) ---
@rag_bp.route('/projets/<int:projet_id>/documents/<int:document_id>/delete', methods=['POST'])
@login_required
def delete_rag_document_route(projet_id, document_id):
    projet = Projet.query.filter_by(id=projet_id, organisation_id=current_user.organisation_id).first_or_404()
    if not projet: # Should be handled by first_or_404
        return jsonify(success=False, error="Projet non trouvé ou accès non autorisé."), 403

    rag_doc_to_delete = RagDocument.query.filter_by(id=document_id, projet_id=projet_id).first_or_404()
    original_filename_to_delete = rag_doc_to_delete.original_filename

    try:
        _internal_delete_rag_document_logic(projet_id, document_id, original_filename_to_delete)
        return jsonify(success=True, message=f"Document '{original_filename_to_delete}' et ses données d'index ont été supprimés avec succès.")
    except ValueError as ve: # Erreur de configuration
        logger.error(f"Erreur de valeur (config) lors de la suppression du document RAG {document_id}: {ve}", exc_info=True)
        return jsonify(success=False, error=str(ve)), 400
    except FileNotFoundError as fnfe:
        logger.error(f"Fichier non trouvé lors de la suppression du document RAG {document_id}: {fnfe}", exc_info=True)
        # Peut-être que le fichier .npz ou original n'existait plus, mais on continue la suppression DB
        # On pourrait retourner une erreur partielle ou un avertissement. Pour l'instant, erreur 500.
        return jsonify(success=False, error=f"Un fichier attendu n'a pas été trouvé: {str(fnfe)}"), 404
    except Exception as e:
        logger.error(f"Erreur inattendue lors de la suppression du document RAG {document_id}: {e}", exc_info=True)
        return jsonify(success=False, error="Erreur serveur lors de la suppression du document."), 500

def _internal_delete_rag_document_logic(projet_id, rag_document_db_id_to_delete, original_filename_to_delete):
    """
    Logique interne pour supprimer un document RAG et ses traces.
    Lève des exceptions en cas d'erreur au lieu de retourner des réponses JSON.
    """
    logger.info(f"Début suppression interne pour Doc ID {rag_document_db_id_to_delete} ('{original_filename_to_delete}') du projet {projet_id}.")

    # 1. Supprimer le fichier original stocké
    physical_file_path_to_delete = get_original_document_save_path(projet_id, original_filename_to_delete)
    if os.path.exists(physical_file_path_to_delete):
        os.remove(physical_file_path_to_delete)
        logger.info(f"Fichier original '{original_filename_to_delete}' supprimé de {physical_file_path_to_delete}")
    else:
        logger.warning(f"Fichier original '{original_filename_to_delete}' non trouvé à {physical_file_path_to_delete} lors de la suppression.")

    # 2. Mettre à jour l'index RAG (.npz) du projet
    rag_index_file_path = get_rag_filepath(projet_id)
    if os.path.exists(rag_index_file_path):
        logger.info(f"Mise à jour du fichier d'index RAG {rag_index_file_path} pour supprimer les chunks du document ID {rag_document_db_id_to_delete}.")
        with np.load(rag_index_file_path, allow_pickle=True) as data:
            if 'texts' not in data or 'embeddings' not in data or 'metadata' not in data:
                logger.warning(f"Fichier RAG {rag_index_file_path} mal formaté ou incomplet. Suppression des chunks impossible, le fichier pourrait être corrompu.")
                # On pourrait choisir de supprimer le .npz s'il est corrompu, ou de le laisser.
                # Pour l'instant, on logue et on continue pour supprimer l'entrée DB.
            else:
                existing_texts = list(data['texts'])
                existing_embeddings = data['embeddings']
                existing_metadata = list(data['metadata'])

                indices_to_keep = [
                    i for i, meta in enumerate(existing_metadata)
                    if isinstance(meta, dict) and meta.get('rag_document_db_id') != rag_document_db_id_to_delete
                ]

                if len(indices_to_keep) < len(existing_texts): # Si des chunks ont été filtrés
                    remaining_texts = [existing_texts[i] for i in indices_to_keep]
                    remaining_embeddings = existing_embeddings[indices_to_keep] if existing_embeddings.size > 0 else np.array([])
                    remaining_metadata = [existing_metadata[i] for i in indices_to_keep]

                    if not remaining_texts: # Si plus aucun texte, supprimer le fichier .npz
                        os.remove(rag_index_file_path)
                        logger.info(f"Plus aucun document dans l'index RAG pour projet {projet_id}. Fichier {rag_index_file_path} supprimé.")
                    else:
                        np.savez_compressed(rag_index_file_path,
                                            texts=np.array(remaining_texts, dtype=object),
                                            embeddings=remaining_embeddings,
                                            metadata=np.array(remaining_metadata, dtype=object))
                        logger.info(f"Index RAG {rag_index_file_path} mis à jour. {len(remaining_texts)} chunks restants.")
                else:
                    logger.info(f"Aucun chunk correspondant au document ID {rag_document_db_id_to_delete} trouvé dans l'index {rag_index_file_path}. Aucune modification de l'index.")
    else:
        logger.warning(f"Fichier d'index RAG {rag_index_file_path} non trouvé lors de la tentative de suppression des chunks du document ID {rag_document_db_id_to_delete}.")

    # 3. Supprimer l'entrée de la base de données
    rag_doc_db_entry = RagDocument.query.get(rag_document_db_id_to_delete)
    if rag_doc_db_entry:
        db.session.delete(rag_doc_db_entry)
        db.session.commit()
        logger.info(f"Entrée RagDocument ID {rag_document_db_id_to_delete} ('{original_filename_to_delete}') supprimée de la base de données.")
    else:
        logger.warning(f"Entrée RagDocument ID {rag_document_db_id_to_delete} non trouvée en BDD lors de la tentative de suppression.")


# --- Route pour visualiser un document RAG original ---
@rag_bp.route('/projets/<int:projet_id>/documents/view/<path:filename>', methods=['GET'])
@login_required
def view_rag_document_original(projet_id, filename):
    projet = Projet.query.filter_by(id=projet_id, organisation_id=current_user.organisation_id).first_or_404()
    if not projet: # Should be handled by first_or_404
        flash("Projet non trouvé ou accès non autorisé.", "danger")
        return redirect(url_for('projets.projets'))

    # Vérifier que le document demandé appartient bien à ce projet et existe en BDD
    rag_doc_entry = RagDocument.query.filter_by(projet_id=projet_id, original_filename=filename).first()
    if not rag_doc_entry:
        logger.warning(f"Tentative de visualisation du document RAG non trouvé en BDD: projet {projet_id}, fichier {filename}")
        flash("Document non trouvé ou non associé à ce projet.", "warning")
        return redirect(url_for('projets.projet_detail', projet_id=projet_id, _anchor='tab5-content'))

    try:
        directory_path = os.path.dirname(get_original_document_save_path(projet_id, filename))
        return send_from_directory(directory_path, filename, as_attachment=False)
    except ValueError as ve: # Erreur de configuration (RAG_UPLOAD_FOLDER manquant)
        logger.error(f"Erreur de configuration pour visualiser le document: {ve}")
        flash("Erreur de configuration serveur.", "danger")
        return redirect(url_for('projets.projet_detail', projet_id=projet_id, _anchor='tab5-content'))
    except FileNotFoundError:
        logger.error(f"Fichier RAG original non trouvé sur le disque: {filename} pour projet {projet_id}")
        flash("Le fichier demandé n'a pas été trouvé sur le serveur.", "warning")
        return redirect(url_for('projets.projet_detail', projet_id=projet_id, _anchor='tab5-content'))
    except Exception as e:
        logger.error(f"Erreur inattendue lors de la visualisation du document {filename}: {e}", exc_info=True)
        flash("Une erreur est survenue lors de la tentative de visualisation du document.", "danger")
        return redirect(url_for('projets.projet_detail', projet_id=projet_id, _anchor='tab5-content'))
